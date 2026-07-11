#!/usr/bin/env python3
"""Post-process the pandoc-generated docx into Information Sciences (Elsevier)
manuscript style: Times New Roman 12pt, double line spacing, continuous line
numbering, centered page numbers, black headings, centered title block.

Run after pandoc (see build_docx.sh). Idempotent.
"""
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

P = sys.argv[1] if len(sys.argv) > 1 else "Fed-CORE_draft.docx"
d = Document(P)

# 1) document default font = Times New Roman 12pt
styles_el = d.styles.element
dd = styles_el.find(qn('w:docDefaults'))
if dd is None:
    dd = OxmlElement('w:docDefaults'); styles_el.insert(0, dd)
rp = dd.find(qn('w:rPrDefault'))
if rp is None:
    rp = OxmlElement('w:rPrDefault'); dd.append(rp)
rPrDef = rp.find(qn('w:rPr'))
if rPrDef is None:
    rPrDef = OxmlElement('w:rPr'); rp.append(rPrDef)
rf = rPrDef.find(qn('w:rFonts'))
if rf is None:
    rf = OxmlElement('w:rFonts'); rPrDef.insert(0, rf)
for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
    rf.set(qn(a), 'Times New Roman')
sz = rPrDef.find(qn('w:sz'))
if sz is None:
    sz = OxmlElement('w:sz'); rPrDef.append(sz)
sz.set(qn('w:val'), '24')  # 12pt


def set_style(nm, size=None, bold=None):
    try:
        st = d.styles[nm]
    except KeyError:
        return
    st.font.name = 'Times New Roman'
    if size:
        st.font.size = Pt(size)
    if bold is not None:
        st.font.bold = bold


for nm in ('Normal', 'Body Text', 'First Paragraph', 'Compact', 'Footer', 'Header'):
    set_style(nm, 12)
for nm, szp in (('Heading 1', 14), ('Heading 2', 13), ('Heading 3', 12), ('Title', 16)):
    set_style(nm, szp, True)

# 2) double line spacing on body
for nm in ('Normal', 'Body Text'):
    try:
        pf = d.styles[nm].paragraph_format
        pf.line_spacing = 2.0
        pf.space_after = Pt(0)
    except KeyError:
        pass

# 3) black headings (strip theme color at style and run level)
for nm in ('Heading 1', 'Heading 2', 'Heading 3', 'Title'):
    try:
        rpr = d.styles[nm].element.get_or_add_rPr()
        for c in rpr.findall(qn('w:color')):
            rpr.remove(c)
        col = OxmlElement('w:color'); col.set(qn('w:val'), '000000'); rpr.append(col)
    except Exception:
        pass


def force_black(run):
    rpr = run._r.get_or_add_rPr()
    for c in rpr.findall(qn('w:color')):
        rpr.remove(c)
    col = OxmlElement('w:color'); col.set(qn('w:val'), '000000'); rpr.append(col)


# 4) center the title block + black heading runs
prefixes = ('Fed-CORE', 'Sanghoon', '[Department', 'Corresponding', 'E-mail')
for i, p in enumerate(d.paragraphs):
    sn = (p.style.name or '')
    if sn.startswith('Heading') or sn == 'Title':
        for r in p.runs:
            force_black(r)
    if i < 8 and (p.text or '').strip().startswith(prefixes):
        p.alignment = AL.CENTER

# 5) qFormat must precede pPr/rPr inside each w:style
pre = (qn('w:pPr'), qn('w:rPr'))
for st in styles_el.findall(qn('w:style')):
    q = st.find(qn('w:qFormat'))
    if q is None:
        continue
    first_pr = next((c for c in list(st) if c.tag in pre), None)
    if first_pr is not None and list(st).index(q) > list(st).index(first_pr):
        st.remove(q); first_pr.addprevious(q)

# 6) continuous line numbering + centered page-number footer
for sec in d.sections:
    sectPr = sec._sectPr
    if sectPr.find(qn('w:lnNumType')) is None:
        ln = OxmlElement('w:lnNumType')
        ln.set(qn('w:countBy'), '1'); ln.set(qn('w:restart'), 'continuous'); ln.set(qn('w:distance'), '454')
        cols = sectPr.find(qn('w:cols'))
        (cols.addprevious(ln) if cols is not None else sectPr.append(ln))
    sec.different_first_page_header_footer = False
    f = sec.footer; f.is_linked_to_previous = False
    fp = f.paragraphs[0] if f.paragraphs else f.add_paragraph()
    fp.text = ''; fp.alignment = AL.CENTER
    fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), ' PAGE ')
    r = OxmlElement('w:r'); tt = OxmlElement('w:t'); tt.text = '1'; r.append(tt); fld.append(r)
    fp._p.append(fld)

# 7) single-space the reference entries (paragraphs starting with "[n] ")
import re as _re
for p in d.paragraphs:
    if _re.match(r'^\[\d+\]\s', (p.text or '')):
        pf = p.paragraph_format
        pf.line_spacing = 1.0
        pf.space_after = Pt(6)


# 8) tables: booktabs look (thick top/bottom rule, thin rule under header,
#    no vertical/inner borders), 10pt Times, single-spaced, tight margins
def _border(tag, val, sz):
    el = OxmlElement(f'w:{tag}')
    el.set(qn('w:val'), val)
    el.set(qn('w:sz'), sz)          # eighths of a point
    el.set(qn('w:space'), '0')
    el.set(qn('w:color'), '000000')
    return el


def _cell_len(cell):
    """Longest visual line in a cell; OMML math text weighted 1.35x since
    math glyphs render wider than the surrounding text."""
    longest = 0.0
    for p in cell.paragraphs:
        n_plain = len(p.text or '')
        n_math = sum(len(mt.text or '')
                     for mt in p._p.findall('.//' + qn('m:t')))
        longest = max(longest, n_plain + 1.35 * n_math)
    return longest


TOTAL_TW = 9360  # ~6.5 in of usable width

for t in d.tables:
    tbl = t._tbl
    tblPr = tbl.tblPr
    # content-proportional fixed column widths
    ncols = max([len(t.columns)] + [len(r.cells) for r in t.rows])
    if ncols == 0:
        continue
    maxlens = []
    for ci in range(ncols):
        m = 0
        for row in t.rows:
            try:
                m = max(m, _cell_len(row.cells[ci]))
            except IndexError:
                pass
        maxlens.append(min(max(m, 6), 34))
    weights = [l ** 0.85 for l in maxlens]
    ws = [max(900, int(TOTAL_TW * w / sum(weights))) for w in weights]
    grid = tbl.find(qn('w:tblGrid'))
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement('w:tblGrid')
    for w in ws:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(w)); grid.append(gc)
    tblPr.addnext(grid)
    for lay in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(lay)
    lay = OxmlElement('w:tblLayout'); lay.set(qn('w:type'), 'fixed')
    tblPr.append(lay)
    for w in tblPr.findall(qn('w:tblW')):
        tblPr.remove(w)
    tw = OxmlElement('w:tblW'); tw.set(qn('w:w'), str(sum(ws))); tw.set(qn('w:type'), 'dxa')
    tblPr.append(tw)
    for row in t.rows:
        for ci, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            for x in tcPr.findall(qn('w:tcW')):
                tcPr.remove(x)
            tcw = OxmlElement('w:tcW')
            tcw.set(qn('w:w'), str(ws[min(ci, ncols - 1)]))
            tcw.set(qn('w:type'), 'dxa')
            tcPr.append(tcw)
    # booktabs borders
    for b in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(b)
    borders = OxmlElement('w:tblBorders')
    borders.append(_border('top', 'single', '12'))
    borders.append(_border('bottom', 'single', '12'))
    for tag in ('left', 'right', 'insideH', 'insideV'):
        borders.append(_border(tag, 'none', '0'))
    tblPr.append(borders)
    # tight default cell margins
    for m in tblPr.findall(qn('w:tblCellMar')):
        tblPr.remove(m)
    mar = OxmlElement('w:tblCellMar')
    for side, wtw in (('top', '20'), ('bottom', '20'), ('left', '80'), ('right', '80')):
        el = OxmlElement(f'w:{side}'); el.set(qn('w:w'), wtw); el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tblPr.append(mar)
    for ri, row in enumerate(t.rows):
        for cell in row.cells:
            if ri == 0:  # header: bold + thin rule below
                tcPr = cell._tc.get_or_add_tcPr()
                for tb in tcPr.findall(qn('w:tcBorders')):
                    tcPr.remove(tb)
                tcB = OxmlElement('w:tcBorders')
                tcB.append(_border('bottom', 'single', '6'))
                tcPr.append(tcB)
            for p in cell.paragraphs:
                pf = p.paragraph_format
                pf.line_spacing = 1.0
                pf.space_before = Pt(1)
                pf.space_after = Pt(1)
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = 'Times New Roman'
                    if ri == 0:
                        r.font.bold = True

# 8b) captions (Figure N. / Table N.) single-spaced 10.5pt
import re as _re
_cap = _re.compile(r'^(Figure|Table)\s+\d+\.')
for p in d.paragraphs:
    if _cap.match(p.text.strip()):
        pf = p.paragraph_format
        pf.line_spacing = 1.0
        pf.space_before = Pt(4)
        pf.space_after = Pt(8)
        for r in p.runs:
            r.font.size = Pt(10.5)

# 9) suppress proofing marks (red/green underlines) in rendered output
try:
    settings = d.settings.element
    for tag in ('w:hideSpellingErrors', 'w:hideGrammaticalErrors'):
        if settings.find(qn(tag)) is None:
            settings.insert(0, OxmlElement(tag))
except Exception:
    pass

d.save(P)
print(f"[ins_format] applied Information Sciences style to {P}")
